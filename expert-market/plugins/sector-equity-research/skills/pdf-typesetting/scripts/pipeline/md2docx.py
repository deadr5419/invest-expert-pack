# -*- coding: utf-8 -*-
"""
md2docx.py — 简版 Markdown → Word(.docx) 转换器（研报结构）
覆盖：封面元信息(#品牌+##标题+>元信息) / 目录(**章节**：说明) / 多级标题 / 表格 / 列表 / 引用 / 粗体
用法: python md2docx.py <input.md> <output.docx>
依赖: python-docx（中文字体用系统字体，天然不方框）
"""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

EA_BODY = "宋体"
EA_HEAD = "黑体"

def set_font(run, name_ea=EA_BODY, size=10.5, bold=False, color=None):
    run.font.name = name_ea
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_ea)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_runs(p, text, size=10.5, bold_base=False, color=None):
    """处理 **粗体** 与 `代码` 内联标记"""
    for part in re.split(r'(\*\*.+?\*\*|`[^`]+`)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            set_font(p.add_run(part[2:-2]), EA_BODY, size, True, color)
        elif part.startswith('`') and part.endswith('`'):
            set_font(p.add_run(part[1:-1]), EA_BODY, size, bold_base, color)
        else:
            set_font(p.add_run(part), EA_BODY, size, bold_base, color)

def add_heading(doc, text, level):
    sizes = {1: 16, 2: 13.5, 3: 12, 4: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(text), EA_HEAD if level <= 1 else EA_BODY,
             sizes.get(level, 11), True, (15, 23, 42) if level <= 2 else None)
    return p

def add_para(doc, text, indent=True, size=10.5):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    add_runs(p, text, size)
    return p

def add_table(doc, rows):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = table.cell(i, j)
            cell.paragraphs[0].text = ''
            txt = row[j] if j < len(row) else ''
            add_runs(cell.paragraphs[0], txt, 9, bold_base=(i == 0))
    doc.add_paragraph()

def parse_table_cells(ln):
    return [c.strip() for c in ln.strip().strip('|').split('|')]

def is_sep_row(cells):
    return all(re.fullmatch(r':?-{2,}:?', c) for c in cells)

def main(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = EA_BODY
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), EA_BODY)
    normal.font.size = Pt(10.5)

    i, n = 0, len(lines)
    state = 'cover'   # cover -> toc -> body
    title_done = False
    brand_cands = []  # 封面 H1 候选（无 ## 主标题时取第一个作大标题）
    # 预扫描：封面段（首个 --- 前）是否有 ## 主标题（非"封面页"类标记）
    cover_has_h2 = False
    for ln in lines[:24]:
        m2 = re.match(r'^##\s+(.+)$', ln.strip())
        if m2 and m2.group(1).strip() not in ('封面', '封面页', '封面信息', '封面元信息'):
            cover_has_h2 = True
            break

    while i < n:
        s = lines[i].strip()

        # 分隔线
        if re.fullmatch(r'-{3,}', s):
            if state == 'cover':
                if not title_done and brand_cands:
                    # 无 ## 主标题（如简版一页纸）：首个 H1 作大标题，其余作品牌行
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(24)
                    set_font(p.add_run(brand_cands[0]), EA_HEAD, 20, True, (15, 23, 42))
                    title_done = True
                    brand_cands = brand_cands[1:]
                for b in brand_cands:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_font(p.add_run(b), EA_BODY, 11, False, (100, 116, 139))
                state = 'toc'
                doc.add_page_break()
            elif state == 'toc':
                state = 'body'
                doc.add_page_break()
            i += 1
            continue

        # 表格（连续 | 行）
        if s.startswith('|') and i + 1 < n and lines[i + 1].strip().startswith('|'):
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                cells = parse_table_cells(lines[i].strip())
                if not is_sep_row(cells):
                    rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue

        # 标题
        m = re.match(r'^(#{1,4})\s+(.*)$', s)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if state == 'cover' and level == 1:
                if not cover_has_h2 and not title_done and not brand_cands:
                    # 无 ## 主标题（如简版一页纸）：首个 H1 直接作大标题
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(24)
                    set_font(p.add_run(text), EA_HEAD, 20, True, (15, 23, 42))
                    title_done = True
                else:
                    # 封面内其余 H1 暂存为品牌/机构候选，分隔线处统一渲染
                    brand_cands.append(text)
            elif state == 'cover' and level == 2 and not title_done:
                if text in ('封面', '封面页', '封面信息', '封面元信息'):
                    pass  # 封面结构标记行，跳过
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(24)
                    set_font(p.add_run(text), EA_HEAD, 20, True, (15, 23, 42))
                    title_done = True
            elif level == 2 and text == '目录':
                state = 'toc'
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(10)
                set_font(p.add_run('目录'), EA_HEAD, 16, True, (15, 23, 42))
            elif state == 'toc':
                pass  # 目录内的其他标题忽略
            else:
                add_heading(doc, text, min(level, 4))
            i += 1
            continue

        # 引用块（封面元信息 / 正文引用）
        if s.startswith('> '):
            t = s[2:].strip()
            if state == 'cover':
                if t.startswith('免责'):
                    p = doc.add_paragraph()
                    set_font(p.add_run('⚠ ' + t), EA_BODY, 8.5, False, (192, 57, 43))
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_font(p.add_run(t.replace('**', '')), EA_BODY, 9.5, False, (100, 116, 139))
            elif state == 'body':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                set_font(p.add_run(t.replace('**', '')), EA_BODY, 9.5, False, (51, 71, 91))
            i += 1
            continue

        # 目录项（**章节**：说明 / - 章节 说明）
        if state == 'toc' and (s.startswith('**') or s.startswith('- ')):
            if s.startswith('**') and '：' in s:
                name, rest = s.split('：', 1)
                name = name.strip('*')
            else:
                t = s.lstrip('- ').strip()
                if '：' in t:
                    name, rest = t.split('：', 1)
                else:
                    name, rest = t, ''
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            set_font(p.add_run(name), EA_BODY, 10.5, True)
            set_font(p.add_run('　' + rest), EA_BODY, 9.5, False, (100, 116, 139))
            i += 1
            continue

        # 无序列表
        if re.match(r'^[-*]\s+', s):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, re.sub(r'^[-*]\s+', '', s), 10.5)
            i += 1
            continue

        # 有序列表
        m = re.match(r'^(\d+)[.、]\s+(.*)$', s)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_runs(p, m.group(2), 10.5)
            i += 1
            continue

        # 普通段落
        if s:
            add_para(doc, s)
        i += 1

    doc.save(docx_path)
    print("DOCX OK:", docx_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        sys.exit("用法: python md2docx.py <input.md> <output.docx>")
    main(sys.argv[1], sys.argv[2])
